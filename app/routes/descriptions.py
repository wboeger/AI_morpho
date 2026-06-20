import io
import os
from datetime import datetime, timezone
from flask import Blueprint, render_template, request, jsonify, flash, redirect, url_for, send_file, current_app
from flask_login import login_required, current_user
from app import db
from app.models import Project, Specimen, Structure, TaxonomicGroup
from app.descriptions import generate_species_description, generate_group_diagnosis

descriptions_bp = Blueprint('descriptions', __name__)


@descriptions_bp.route('/project/<int:project_id>/descriptions')
@login_required
def species_descriptions(project_id):
    project = Project.query.get_or_404(project_id)
    specimens = Specimen.query.filter_by(project_id=project_id).order_by(Specimen.species_name).all()
    return render_template('descriptions/species_list.html',
                           project=project, specimens=specimens)


def _specimen_images(specimen_id):
    """Return list of {structure_type, image_url, abs_path} for all structures with images."""
    ST_ORDER = ['hook', 'anchor', 'superficial_bar', 'deep_bar', 'mco']
    structs = Structure.query.filter_by(specimen_id=specimen_id).all()
    seen = {}
    for s in structs:
        if s.image_path and s.structure_type not in seen:
            seen[s.structure_type] = s.image_path
    return [
        {
            'structure_type': st,
            'label': st.replace('_', ' ').title(),
            'image_url': '/uploads/' + seen[st],
            'abs_path': os.path.join(current_app.config['UPLOAD_FOLDER'], seen[st]),
        }
        for st in ST_ORDER if st in seen
    ] + [
        {
            'structure_type': st,
            'label': st.replace('_', ' ').title(),
            'image_url': '/uploads/' + path,
            'abs_path': os.path.join(current_app.config['UPLOAD_FOLDER'], path),
        }
        for st, path in seen.items() if st not in ST_ORDER
    ]


@descriptions_bp.route('/project/<int:project_id>/descriptions/<int:specimen_id>')
@login_required
def view_description(project_id, specimen_id):
    project = Project.query.get_or_404(project_id)
    specimen = Specimen.query.get_or_404(specimen_id)
    description = generate_species_description(specimen_id, project_id)
    images = _specimen_images(specimen_id)
    return render_template('descriptions/view_description.html',
                           project=project, specimen=specimen,
                           description=description, images=images)


@descriptions_bp.route('/api/project/<int:project_id>/descriptions/<int:specimen_id>/regenerate', methods=['POST'])
@login_required
def regenerate_description(project_id, specimen_id):
    description = generate_species_description(specimen_id, project_id)
    return jsonify({'description': description})


@descriptions_bp.route('/project/<int:project_id>/descriptions/<int:specimen_id>/docx')
@login_required
def export_description_docx(project_id, specimen_id):
    from docx import Document
    from docx.shared import Inches, Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    project  = Project.query.get_or_404(project_id)
    specimen = Specimen.query.get_or_404(specimen_id)
    description = generate_species_description(specimen_id, project_id)
    images = _specimen_images(specimen_id)

    doc = Document()

    # ── Page margins ──────────────────────────────────────────────
    section = doc.sections[0]
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(3)

    # ── Species heading (italic) ──────────────────────────────────
    h = doc.add_heading('', level=1)
    run = h.add_run(specimen.species_name)
    run.italic = True
    h.paragraph_format.space_after = Pt(6)

    # ── Image gallery ─────────────────────────────────────────────
    valid_images = [img for img in images if os.path.exists(img['abs_path'])]
    if valid_images:
        gallery_heading = doc.add_heading('Illustrations', level=2)
        gallery_heading.paragraph_format.space_after = Pt(4)

        cols = min(len(valid_images), 3)
        table = doc.add_table(rows=0, cols=cols)
        table.style = 'Table Grid'

        # fill rows
        for row_start in range(0, len(valid_images), cols):
            row = table.add_row()
            for ci, img in enumerate(valid_images[row_start:row_start + cols]):
                cell = row.cells[ci]
                # image
                img_p = cell.paragraphs[0]
                img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                img_run = img_p.add_run()
                try:
                    img_run.add_picture(img['abs_path'], width=Inches(1.7))
                except Exception:
                    img_p.text = img['label']
                # caption below
                cap_p = cell.add_paragraph(img['label'])
                cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if cap_p.runs:
                    cap_p.runs[0].font.size = Pt(8)
                    cap_p.runs[0].italic = True

        doc.add_paragraph()   # spacing after table

    # ── Description body ──────────────────────────────────────────
    doc.add_heading('Description', level=2)
    # Each \n\n-separated block becomes its own paragraph; skip the first
    # line (species name is already in the heading).
    blocks = description.split('\n\n')
    for block in blocks:
        block = block.strip()
        if not block or block == specimen.species_name:
            continue
        p = doc.add_paragraph(block)
        p.paragraph_format.space_after = Pt(6)
        if p.runs:
            p.runs[0].font.size = Pt(11)

    # ── Stream to browser ─────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    safe = specimen.species_name.replace(' ', '_')
    return send_file(
        buf,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name=f'{safe}_description.docx',
    )


@descriptions_bp.route('/project/<int:project_id>/diagnoses')
@login_required
def diagnoses(project_id):
    project = Project.query.get_or_404(project_id)
    groups = TaxonomicGroup.query.filter_by(project_id=project_id).order_by(TaxonomicGroup.rank, TaxonomicGroup.name).all()
    return render_template('descriptions/diagnoses.html', project=project, groups=groups)


@descriptions_bp.route('/project/<int:project_id>/diagnoses/new', methods=['GET', 'POST'])
@login_required
def new_group(project_id):
    project = Project.query.get_or_404(project_id)
    specimens = Specimen.query.filter_by(project_id=project_id).order_by(Specimen.species_name).all()
    all_species = sorted(set(s.species_name for s in specimens))

    if request.method == 'POST':
        name = request.form.get('name', '').strip()
        rank = request.form.get('rank', '').strip()
        species_list = request.form.getlist('species')

        if not name:
            flash('Group name is required.', 'error')
        else:
            group = TaxonomicGroup(
                project_id=project_id,
                name=name,
                rank=rank,
                included_species=species_list,
            )
            db.session.add(group)
            db.session.commit()

            # Auto-generate diagnosis
            diagnosis = generate_group_diagnosis(group.id, project_id)
            group.diagnosis_text = diagnosis
            group.diagnosis_generated_at = datetime.now(timezone.utc)
            db.session.commit()

            flash(f'Group "{name}" created with diagnosis.', 'success')
            return redirect(url_for('descriptions.diagnoses', project_id=project_id))

    return render_template('descriptions/new_group.html',
                           project=project, all_species=all_species)


@descriptions_bp.route('/project/<int:project_id>/diagnoses/<int:group_id>')
@login_required
def view_diagnosis(project_id, group_id):
    project = Project.query.get_or_404(project_id)
    group = TaxonomicGroup.query.get_or_404(group_id)
    return render_template('descriptions/view_diagnosis.html',
                           project=project, group=group)


@descriptions_bp.route('/api/project/<int:project_id>/diagnoses/<int:group_id>/regenerate', methods=['POST'])
@login_required
def regenerate_diagnosis(project_id, group_id):
    group = TaxonomicGroup.query.get_or_404(group_id)
    diagnosis = generate_group_diagnosis(group_id, project_id)
    group.diagnosis_text = diagnosis
    group.diagnosis_generated_at = datetime.now(timezone.utc)
    db.session.commit()
    return jsonify({'diagnosis': diagnosis})


@descriptions_bp.route('/api/project/<int:project_id>/diagnoses/<int:group_id>/save', methods=['POST'])
@login_required
def save_diagnosis(project_id, group_id):
    group = TaxonomicGroup.query.get_or_404(group_id)
    data = request.get_json()
    group.diagnosis_text = data.get('text', '')
    db.session.commit()
    return jsonify({'status': 'ok'})
